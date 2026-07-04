import json
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .config import settings


PROJECT_ROOT = Path(__file__).resolve().parents[2]
TEMPLATES_DIR = PROJECT_ROOT / "templates"
REPORT_TEMPLATE_REGISTRY = TEMPLATES_DIR / "report_templates.json"
REPORT_TYPE_NAMES = {
    "weekly": "周报",
    "monthly": "月报",
    "internship_summary": "实习总结",
}


def _copy_run_format(paragraph, like) -> None:
    if not like or not paragraph.runs or not like.runs:
        return
    src = like.runs[0].font
    dst = paragraph.runs[0].font
    dst.name = src.name
    dst.size = src.size
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline


def _set_text(paragraph, text: str, like=None) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
    _copy_run_format(paragraph, like)


def _insert_after(paragraph, text: str, like=None):
    new = OxmlElement("w:p")
    paragraph._p.addnext(new)
    p = Paragraph(new, paragraph._parent)
    p.style = paragraph.style
    if like is not None:
        p.style = like.style
    _set_text(p, text, like)
    return p


def _replace(paragraph, values: Dict[str, str]) -> bool:
    text = paragraph.text
    if not any(k in text for k in values):
        return False
    for k, v in values.items():
        text = text.replace(k, v)
    _set_text(paragraph, text)
    return True


def _delete(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _iter_paragraphs(doc) -> Iterable[Paragraph]:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _values(values: Dict[str, Any]) -> Dict[str, str]:
    out = {}
    for k, v in values.items():
        text = "" if v is None else str(v)
        out[k] = text
        if not (k.startswith("{{") and k.endswith("}}")):
            out[f"{{{{{k}}}}}"] = text
    return out


def load_report_templates(path: Path = REPORT_TEMPLATE_REGISTRY) -> List[Dict[str, Any]]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)["templates"]


def get_report_template(report_type: str, template_id: str = "") -> Dict[str, Any]:
    for item in load_report_templates():
        if template_id and item["id"] == template_id:
            return item
        if not template_id and item["type"] == report_type:
            return item
    raise ValueError(f"report template not found: {template_id or report_type}")


def render_docx(
    date: str,
    title: str,
    body: List[str],
    values: Dict[str, str],
    out: Path,
    template_path: Path = None,
) -> Path:
    doc = Document(str(template_path or settings().template_path))
    template_samples = [
        p for p in doc.paragraphs
        if p.text in {"标题1", "标题2", "标题3", "标题4"} or p.text.startswith("这是正文的示例")
    ]
    title_like = next((p for p in doc.paragraphs if p.text == "标题1"), None)
    body_like = next((p for p in doc.paragraphs if p.text.startswith("这是正文的示例")), None)
    merged = _values(values)
    name = merged.get("{{姓名}}", "")
    klass = merged.get("{{班级}}", "")
    student_id = merged.get("{{学号}}", "")
    merged.update(
        {
            "{{日期}}": date,
            "{{标题}}": title,
            "{{日记标题}}": title,
            "{{正文}}": "\n".join(body),
            "张三": name,
            "微电子XXXX": klass,
            "XXXXXXXXXX": student_id,
            "XXXX年XX月XX日": date,
        }
    )
    used_title_placeholder = used_body_placeholder = False
    for p in _iter_paragraphs(doc):
        before = p.text
        if "{{日记标题}}" in before or "{{标题}}" in before:
            _set_text(p, before.replace("{{日记标题}}", title).replace("{{标题}}", title), title_like)
            used_title_placeholder = True
            continue
        if "{{正文}}" in before:
            paragraphs = body or [""]
            _set_text(p, before.replace("{{正文}}", paragraphs[0]), body_like)
            anchor = p
            for text in paragraphs[1:]:
                anchor = _insert_after(anchor, text, body_like)
            used_body_placeholder = True
            continue
        changed = _replace(p, merged)
        used_title_placeholder = used_title_placeholder or "{{日记标题}}" in before
        used_body_placeholder = used_body_placeholder or "{{正文}}" in before
        if changed:
            continue
        if before == "标题1":
            _set_text(p, title, title_like)
            used_title_placeholder = True
        elif before.startswith("这是正文的示例") and not used_body_placeholder:
            _set_text(p, "\n".join(body), body_like)
            used_body_placeholder = True
    if used_title_placeholder and used_body_placeholder:
        for p in list(template_samples):
            if p.text not in {"标题1", "标题2", "标题3", "标题4"} and not p.text.startswith("这是正文的示例"):
                continue
            _delete(p)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def render_report_docx(
    report_type: str,
    start_date: str,
    end_date: str,
    title: str,
    body: List[str],
    values: Dict[str, str],
    out: Path,
    template_id: str = "",
) -> Path:
    template = get_report_template(report_type, template_id)
    report_values = _values(values)
    report_values.update(
        {
            "{{报告类型}}": REPORT_TYPE_NAMES.get(report_type, report_type),
            "{{开始日期}}": start_date,
            "{{结束日期}}": end_date,
            "专业实习 实习日记": f"专业实习 {REPORT_TYPE_NAMES.get(report_type, report_type)}",
        }
    )
    name = title or f"{report_values['{{报告类型}}']}（{start_date} 至 {end_date}）"
    body_lines = body if isinstance(body, list) else str(body).splitlines()
    return render_docx(
        end_date,
        name,
        body_lines,
        report_values,
        out,
        TEMPLATES_DIR / template.get("file", "diary_template_working.docx"),
    )
