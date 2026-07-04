from pathlib import Path

from docx import Document

from app.config import settings
from app.word_renderer import render_docx


def make_template(path: Path):
    d = Document()
    d.add_paragraph("姓名：{{姓名}}")
    d.add_paragraph("{{日记标题}}")
    d.add_paragraph("{{正文}}")
    d.save(str(path))


def test_render_docx(tmp_path, monkeypatch):
    tpl = tmp_path / "tpl.docx"
    out = tmp_path / "out.docx"
    make_template(tpl)
    monkeypatch.setenv("TEMPLATE_PATH", str(tpl))
    settings.cache_clear()
    render_docx("2026-07-01", "专业实习日记", ["第一段", "第二段"], {"{{姓名}}": "张三"}, out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "张三" in text
    assert "专业实习日记" in text
    assert "第一段" in text


def test_render_official_template_with_profile(tmp_path, monkeypatch):
    tpl = Path(__file__).resolve().parents[2] / "2026年专业实习 实习报告和日记模板" / "2026年专业实习 实习日记模板.docx"
    if not tpl.exists():
        return
    out = tmp_path / "official.docx"
    monkeypatch.setenv("TEMPLATE_PATH", str(tpl))
    settings.cache_clear()
    render_docx(
        "2026-07-02",
        "专业实习日记",
        ["第一段", "第二段"],
        {"{{姓名}}": "孙梓越", "{{班级}}": "微电子2304", "{{学号}}": "2236515806"},
        out,
    )
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "孙梓越" in text
    assert "微电子2304" in text
    assert "2236515806" in text
    assert "第一段" in text
    assert "这是正文的示例" not in text
    for title in ("标题1", "标题2", "标题3", "标题4"):
        assert title not in text


def test_render_official_template_without_profile_has_no_personal_info(tmp_path, monkeypatch):
    tpl = Path(__file__).resolve().parents[2] / "2026年专业实习 实习报告和日记模板" / "2026年专业实习 实习日记模板.docx"
    if not tpl.exists():
        return
    out = tmp_path / "blank.docx"
    monkeypatch.setenv("TEMPLATE_PATH", str(tpl))
    settings.cache_clear()
    render_docx("2026-07-02", "专业实习日记", ["第一段"], {"{{姓名}}": "", "{{班级}}": "", "{{学号}}": ""}, out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "孙梓越" not in text
    assert "微电子2304" not in text
    assert "2236515806" not in text
    assert "张三" not in text
    assert "微电子XXXX" not in text
    assert "XXXXXXXXXX" not in text
