from fastapi.testclient import TestClient
import sys

from app.config import settings
from app.main import app

H = {"Authorization": "Bearer dev-token"}


def _fake_codex(tmp_path, output):
    p = tmp_path / "fake_codex.py"
    p.write_text(
        "import sys\n"
        "out = sys.argv[sys.argv.index('--output-last-message') + 1]\n"
        f"open(out, 'w', encoding='utf-8').write({output!r})\n",
        encoding="utf-8",
    )
    return f"{sys.executable} {p}"


def test_sort_and_generate_draft(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("TEMPLATE_PATH", str(tmp_path / "tpl.docx"))
    settings.cache_clear()
    from docx import Document

    d = Document()
    d.add_paragraph("姓名：{{姓名}}")
    d.add_paragraph("{{日记标题}}")
    d.add_paragraph("{{正文}}")
    d.save(str(tmp_path / "tpl.docx"))

    c = TestClient(app)
    c.post("/api/entries/text", headers=H, json={"date": "2026-07-01", "content": "学习巡检流程"})
    assert c.post("/api/actions/sort-day", headers=H, json={"date": "2026-07-01"}).status_code == 200
    r = c.post("/api/actions/generate-diary", headers=H, json={"date": "2026-07-01", "word_count": 800})
    assert r.status_code == 200
    assert (tmp_path / "workdays" / "2026-07-01" / "diary_draft.md").exists()


def test_sort_day_uses_codex_cli(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("CODEX_ENABLED", "1")
    monkeypatch.setenv("CODEX_COMMAND", _fake_codex(tmp_path, "# codex sorted\n"))
    settings.cache_clear()

    c = TestClient(app)
    c.post("/api/entries/text", headers=H, json={"date": "2026-07-01", "content": "学习巡检流程"})
    r = c.post("/api/actions/sort-day", headers=H, json={"date": "2026-07-01"})
    assert r.status_code == 200
    out = tmp_path / "workdays" / "2026-07-01" / "sorted_notes.md"
    assert out.read_text(encoding="utf-8") == "# codex sorted\n"


def test_generate_diary_uses_codex_cli(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("TEMPLATE_PATH", str(tmp_path / "tpl.docx"))
    monkeypatch.setenv("CODEX_ENABLED", "1")
    monkeypatch.setenv(
        "CODEX_COMMAND",
        _fake_codex(
            tmp_path,
            '{"date":"2026-07-01","title":"Codex 日记","body_paragraphs":["正文"],"safety_notes":[]}',
        ),
    )
    settings.cache_clear()
    from docx import Document

    d = Document()
    d.add_paragraph("{{日记标题}}")
    d.add_paragraph("{{正文}}")
    d.save(str(tmp_path / "tpl.docx"))

    c = TestClient(app)
    c.post("/api/entries/text", headers=H, json={"date": "2026-07-01", "content": "学习巡检流程"})
    r = c.post("/api/actions/generate-diary", headers=H, json={"date": "2026-07-01", "word_count": 800})
    assert r.status_code == 200
    assert "Codex 日记" in r.json()["draft"]
