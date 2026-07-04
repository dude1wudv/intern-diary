"""One-off script to ensure placeholder paragraphs exist in the working
Word template. Preserves the original template's cover/content and appends
the placeholder block at the end, instead of editing the document visually.

Usage:
    python tools/build_template.py [path/to/diary_template_working.docx]

Default path (relative to this file): ../../templates/diary_template_working.docx
"""
import sys
from pathlib import Path

from docx import Document

DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent.parent / "templates" / "diary_template_working.docx"

PLACEHOLDER_LINES = [
    "姓名：{{姓名}}",
    "班级：{{班级}}",
    "学号：{{学号}}",
    "日期：{{日期}}",
    "",
    "{{日记标题}}",
    "",
    "{{正文}}",
]


def ensure_placeholders(path: Path) -> None:
    doc = Document(str(path))
    existing_text = "\n".join(p.text for p in doc.paragraphs)
    if "{{正文}}" in existing_text and "{{姓名}}" in existing_text:
        print(f"placeholders already present in {path}, no change made")
        return
    for line in PLACEHOLDER_LINES:
        doc.add_paragraph(line)
    doc.save(str(path))
    print(f"appended placeholder block to {path}")


def main() -> None:
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_TEMPLATE
    if not target.exists():
        raise SystemExit(f"template not found: {target}")
    ensure_placeholders(target)


if __name__ == "__main__":
    main()
